from __future__ import annotations

from app.models.resume import ResumeSchema

_PRESETS: dict[str, dict[str, str]] = {
    "tech":       {"font_name": "Arial",            "accent": "#2563eb"},
    "finance":    {"font_name": "Georgia",           "accent": "#1e3a5f"},
    "creative":   {"font_name": "Helvetica Neue",    "accent": "#7c3aed"},
    "healthcare": {"font_name": "Times New Roman",   "accent": "#065f46"},
    "general":    {"font_name": "Helvetica",         "accent": "#374151"},
}

# ReportLab uses its own font names; map preset names to built-in Type-1 faces
# (embedded in every PDF — no system fonts required)
_RL_FONTS: dict[str, tuple[str, str]] = {
    "Arial":          ("Helvetica",   "Helvetica-Bold"),
    "Georgia":        ("Times-Roman", "Times-Bold"),
    "Helvetica Neue": ("Helvetica",   "Helvetica-Bold"),
    "Times New Roman":("Times-Roman", "Times-Bold"),
    "Helvetica":      ("Helvetica",   "Helvetica-Bold"),
}


def _x(text: str | None) -> str:
    """Escape XML entities for ReportLab Paragraph markup."""
    if not text:
        return ""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def generate_pdf(resume: ResumeSchema, industry: str = "general") -> bytes:
    from io import BytesIO
    from reportlab.lib.colors import HexColor, black
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    p = _PRESETS.get(industry, _PRESETS["general"])
    accent = HexColor(p["accent"])
    body_font, bold_font = _RL_FONTS.get(p["font_name"], ("Helvetica", "Helvetica-Bold"))

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )

    gray55 = HexColor("#555555")
    gray44 = HexColor("#444444")

    s_name     = ParagraphStyle("name",     fontName=bold_font,  fontSize=20, leading=24, textColor=accent,  alignment=TA_CENTER, spaceAfter=3)
    s_contact  = ParagraphStyle("contact",  fontName=body_font,  fontSize=9,  leading=12, textColor=gray55,  alignment=TA_CENTER, spaceAfter=6)
    s_section  = ParagraphStyle("section",  fontName=bold_font,  fontSize=10, leading=14, textColor=accent,  spaceBefore=10,      spaceAfter=3)
    s_body     = ParagraphStyle("body",     fontName=body_font,  fontSize=10, leading=14, textColor=black,   spaceAfter=2)
    s_subtitle = ParagraphStyle("subtitle", fontName=body_font,  fontSize=9,  leading=12, textColor=gray44,  spaceAfter=2)
    s_bullet   = ParagraphStyle("bullet",   fontName=body_font,  fontSize=10, leading=14, textColor=black,   leftIndent=14, spaceAfter=1)

    story: list = []
    m = resume.metadata

    # ── Header ─────────────────────────────────────────────────────────────────
    story.append(Paragraph(_x(m.name) or "Name", s_name))
    contacts = [c for c in [m.email, m.phone, m.location, m.linkedin, m.github] if c]
    if contacts:
        story.append(Paragraph("  ·  ".join(_x(c) for c in contacts), s_contact))
    story.append(HRFlowable(width="100%", thickness=1.5, color=accent, spaceAfter=4))

    def _section(title: str) -> None:
        story.append(Paragraph(title, s_section))
        story.append(HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=4))

    def _title_row(left: str, right: str) -> None:
        tbl = Table(
            [[Paragraph(left, s_body), Paragraph(_x(right), s_subtitle)]],
            colWidths=["*", 1.35 * inch],
        )
        tbl.setStyle(TableStyle([
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("ALIGN",         (1, 0), (1, 0),   "RIGHT"),
            ("LEFTPADDING",   (0, 0), (-1, -1), 0),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
            ("TOPPADDING",    (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        story.append(tbl)

    # ── Summary ─────────────────────────────────────────────────────────────────
    if resume.summary:
        _section("SUMMARY")
        story.append(Paragraph(_x(resume.summary), s_body))

    # ── Experience ──────────────────────────────────────────────────────────────
    if resume.experience:
        _section("EXPERIENCE")
        for exp in resume.experience:
            date = exp.start_date + (" – Present" if not exp.end_date else f" – {exp.end_date}")
            _title_row(f"<b>{_x(exp.title)}</b>", date)
            company = _x(exp.company) + (f"  ·  {_x(exp.location)}" if exp.location else "")
            story.append(Paragraph(company, s_subtitle))
            for b in exp.bullets:
                if b.strip():
                    story.append(Paragraph(f"• {_x(b)}", s_bullet))
            story.append(Spacer(1, 4))

    # ── Education ───────────────────────────────────────────────────────────────
    if resume.education:
        _section("EDUCATION")
        for edu in resume.education:
            if edu.start_date and edu.end_date:
                date = f"{edu.start_date} – {edu.end_date}"
            elif edu.end_date:
                date = edu.end_date
            else:
                date = ""
            _title_row(f"<b>{_x(edu.school)}</b>", date)
            degree = _x(edu.degree) + (f" in {_x(edu.field)}" if edu.field else "")
            if edu.gpa:
                degree += f"  ·  GPA: {_x(edu.gpa)}"
            if degree:
                story.append(Paragraph(degree, s_subtitle))
            if edu.honors:
                story.append(Paragraph(_x(edu.honors), s_subtitle))
            story.append(Spacer(1, 4))

    # ── Projects ────────────────────────────────────────────────────────────────
    if resume.projects:
        _section("PROJECTS")
        for proj in resume.projects:
            name_part = _x(proj.name)
            if proj.technologies:
                name_part += f"  ·  {_x(', '.join(proj.technologies))}"
            _title_row(f"<b>{name_part}</b>", proj.url or "")
            if proj.description:
                story.append(Paragraph(_x(proj.description), s_subtitle))
            for b in proj.bullets:
                if b.strip():
                    story.append(Paragraph(f"• {_x(b)}", s_bullet))
            story.append(Spacer(1, 4))

    # ── Skills ──────────────────────────────────────────────────────────────────
    if resume.skills:
        _section("SKILLS")
        for sg in resume.skills:
            items = _x(", ".join(sg.items))
            line = f"<b>{_x(sg.category)}:</b> {items}" if sg.category else items
            story.append(Paragraph(line, s_body))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def generate_cover_letter_pdf(content: str, company_name: str) -> bytes:
    from datetime import date
    from io import BytesIO
    from reportlab.lib.colors import HexColor, black
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=1.0 * inch,
        rightMargin=1.0 * inch,
        topMargin=1.0 * inch,
        bottomMargin=1.0 * inch,
    )

    gray = HexColor("#444444")
    s_date = ParagraphStyle("date",  fontName="Helvetica", fontSize=10, leading=14, textColor=gray, spaceAfter=16)
    s_body = ParagraphStyle("body",  fontName="Helvetica", fontSize=11, leading=16, textColor=black, spaceAfter=12)

    story: list = []
    story.append(Paragraph(date.today().strftime("%B %d, %Y"), s_date))

    for para in content.strip().split("\n\n"):
        para = para.strip()
        if para:
            story.append(Paragraph(_x(para), s_body))

    story.append(Spacer(1, 8))
    story.append(Paragraph("Best regards,", s_body))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def generate_cover_letter_docx(content: str, company_name: str) -> bytes:
    from datetime import date
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    date_p = doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

    for para in content.strip().split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    closing = doc.add_paragraph("Best regards,")
    closing.runs[0].font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_docx(resume: ResumeSchema, industry: str = "general") -> bytes:
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    preset = _PRESETS.get(industry, _PRESETS["general"])
    font_name = preset["font_name"]
    accent_hex = preset["accent"].lstrip("#")
    accent_rgb = RGBColor(int(accent_hex[0:2], 16), int(accent_hex[2:4], 16), int(accent_hex[4:6], 16))

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    m = resume.metadata

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(m.name or "")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.name = font_name

    contact_parts = [p for p in [m.email, m.phone, m.location, m.linkedin, m.github] if p]
    if contact_parts:
        cp = doc.add_paragraph(" · ".join(contact_parts))
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def _heading(title: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = font_name
        run.font.color.rgb = accent_rgb
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), accent_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    if resume.summary:
        _heading("Summary")
        doc.add_paragraph(resume.summary)

    if resume.experience:
        _heading("Experience")
        for exp in resume.experience:
            date_range = exp.start_date + (f" – {exp.end_date}" if exp.end_date else " – Present")
            p = doc.add_paragraph()
            r1 = p.add_run(exp.title)
            r1.bold = True
            r2 = p.add_run(f"  {date_range}")
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            company_line = exp.company + (f" · {exp.location}" if exp.location else "")
            cp = doc.add_paragraph(company_line)
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            for bullet in exp.bullets:
                if bullet.strip():
                    doc.add_paragraph(bullet, style="List Bullet")

    if resume.education:
        _heading("Education")
        for edu in resume.education:
            if edu.start_date and edu.end_date:
                date_range = f"{edu.start_date} – {edu.end_date}"
            elif edu.end_date:
                date_range = edu.end_date
            else:
                date_range = ""
            p = doc.add_paragraph()
            p.add_run(edu.school).bold = True
            if date_range:
                dr = p.add_run(f"  {date_range}")
                dr.font.size = Pt(9)
                dr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            degree_line = edu.degree + (f" in {edu.field}" if edu.field else "")
            if edu.gpa:
                degree_line += f" · GPA: {edu.gpa}"
            if degree_line:
                dp = doc.add_paragraph(degree_line)
                dp.runs[0].font.size = Pt(9)

    if resume.skills:
        _heading("Skills")
        for sg in resume.skills:
            p = doc.add_paragraph()
            p.add_run(f"{sg.category}: ").bold = True
            p.add_run(", ".join(sg.items))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
